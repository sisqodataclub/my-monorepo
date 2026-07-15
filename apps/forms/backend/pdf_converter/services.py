# pdf_converter/services.py
import io
import pdfkit
import os
import tempfile
from django.template import Template, Context
from django.template.loader import render_to_string
from django.conf import settings

# For file conversions
from docx import Document as DocxDocument
import markdown

class PDFService:
    """Core PDF generation engine using pdfkit (wkhtmltopdf)."""

    @staticmethod
    def render_html_to_pdf(html_source, context=None, css_string=None, options=None):
        """
        Convert HTML + CSS + Context data into a PDF using pdfkit.

        Args:
            html_source (str): Raw HTML or template path (e.g., 'pdf/invoice.html').
            context (dict): Data to inject into the HTML.
            css_string (str): Optional custom CSS (injected into <style> tag).
            options (dict): Additional pdfkit options (e.g., page size, margin).

        Returns:
            bytes: The generated PDF file as bytes.
        """
        context = context or {}

        # 1. Render the HTML
        if isinstance(html_source, str) and html_source.endswith('.html'):
            rendered_html = render_to_string(html_source, context)
        else:
            template = Template(html_source)
            rendered_html = template.render(Context(context))

        # 2. Inject custom CSS if provided
        if css_string:
            style_tag = f"<style>{css_string}</style>"
            if '<head>' in rendered_html:
                rendered_html = rendered_html.replace('<head>', f'<head>{style_tag}')
            else:
                rendered_html = f'<html><head>{style_tag}</head><body>{rendered_html}</body></html>'

        # 3. Prepare pdfkit options
        wkhtmltopdf_options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'enable-local-file-access': None,
        }
        if options:
            wkhtmltopdf_options.update(options)

        # 4. Generate PDF using pdfkit
        config = pdfkit.configuration(wkhtmltopdf=settings.WKHTMLTOPDF_PATH) if hasattr(settings, 'WKHTMLTOPDF_PATH') else None
        pdf_bytes = pdfkit.from_string(rendered_html, False, options=wkhtmltopdf_options, configuration=config)
        return pdf_bytes

    @staticmethod
    def generate_from_template(template_obj, context, options=None):
        """Generate a PDF using a stored PDFTemplate model instance."""
        html_content = template_obj.html_content
        css_content = template_obj.css_content
        return PDFService.render_html_to_pdf(html_content, context, css_content, options)

    @staticmethod
    def merge_pdfs(pdf_bytes_list):
        """Merge multiple PDF byte streams into a single PDF file."""
        if len(pdf_bytes_list) == 1:
            return pdf_bytes_list[0]

        import PyPDF2
        merger = PyPDF2.PdfMerger()
        for pdf_bytes in pdf_bytes_list:
            merger.append(io.BytesIO(pdf_bytes))

        output = io.BytesIO()
        merger.write(output)
        output.seek(0)
        return output.getvalue()

    # ============================================================
    # 🆕 Document to HTML Conversion
    # ============================================================

    @staticmethod
    def convert_docx_to_html(file_bytes):
        """Convert a .docx file (bytes) to HTML."""
        doc = DocxDocument(io.BytesIO(file_bytes))
        html_parts = []

        if doc.core_properties.title:
            html_parts.append(f"<h1>{doc.core_properties.title}</h1>")

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.style.name and paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading', '').strip()
                    if level.isdigit():
                        html_parts.append(f"<h{level}>{paragraph.text}</h{level}>")
                    else:
                        html_parts.append(f"<h2>{paragraph.text}</h2>")
                else:
                    html_parts.append(f"<p>{paragraph.text}</p>")

        for table in doc.tables:
            html_parts.append('<table border="1" cellpadding="5" style="border-collapse: collapse; width: 100%;">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')

        return f"""
        <!DOCTYPE html>
        <html>
        <head><meta charset="UTF-8"><title>{doc.core_properties.title or 'Document'}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
            h1, h2, h3, h4, h5, h6 {{ color: #333; }}
            p {{ margin: 10px 0; }}
            table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
            td, th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
        </style></head>
        <body>{''.join(html_parts)}</body></html>
        """

    @staticmethod
    def convert_markdown_to_html(file_bytes):
        """Convert a Markdown file (bytes) to HTML."""
        content = file_bytes.decode('utf-8')
        html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
        return f"""
        <!DOCTYPE html>
        <html>
        <head><meta charset="UTF-8"><title>Markdown Document</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
            h1, h2, h3, h4, h5, h6 {{ color: #333; }}
            code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 4px; }}
            pre {{ background: #f4f4f4; padding: 10px; border-radius: 4px; overflow-x: auto; }}
            blockquote {{ border-left: 4px solid #ddd; padding-left: 20px; color: #666; }}
            table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
            td, th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
        </style></head>
        <body>{html_content}</body></html>
        """

    @staticmethod
    def convert_text_to_html(file_bytes):
        """Convert a plain text file to HTML."""
        content = file_bytes.decode('utf-8')
        paragraphs = [f"<p>{p.strip()}</p>" for p in content.split('\n\n') if p.strip()]
        return f"""
        <!DOCTYPE html>
        <html>
        <head><meta charset="UTF-8"><title>Text Document</title>
        <style>body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }} p {{ margin: 10px 0; }}</style>
        </head>
        <body>{''.join(paragraphs)}</body></html>
        """

    @staticmethod
    def convert_html_to_html(file_bytes):
        """Return HTML content unchanged (just decode)."""
        return file_bytes.decode('utf-8')

    @staticmethod
    def convert_file_to_html(file_bytes, file_extension):
        """Route file to the appropriate converter based on extension."""
        extension = file_extension.lower()

        if extension in ['.docx', '.doc']:
            return PDFService.convert_docx_to_html(file_bytes)
        elif extension in ['.md', '.markdown']:
            return PDFService.convert_markdown_to_html(file_bytes)
        elif extension in ['.txt', '.text']:
            return PDFService.convert_text_to_html(file_bytes)
        elif extension in ['.html', '.htm']:
            return PDFService.convert_html_to_html(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {extension}. Supported: .docx, .md, .txt, .html")
